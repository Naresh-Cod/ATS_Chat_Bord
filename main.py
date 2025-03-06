import streamlit as st
import os
import io
import base64
from PIL import Image
import pdf2image
import google.generativeai as genai
from docx import Document
from fpdf import FPDF
from dotenv import load_dotenv
import requests

# Set page config FIRST - NO CODE OR COMMENTS ABOVE THIS LINE
st.set_page_config(page_title="ATS Resume Expert & Interview Question Generator", layout="wide")

# Load environment variables
load_dotenv()

# Configure Google Gemini API
API_KEY = os.getenv("GOOGLE_API_KEY")
if not API_KEY:
    st.error("GOOGLE_API_KEY not found. Please set it in your environment variables.")
    st.stop()

genai.configure(api_key=API_KEY)

# Configure JSearch API
JSEARCH_API_KEY = os.getenv("JSEARCH_API_KEY")
if not JSEARCH_API_KEY:
    st.error("JSEARCH_API_KEY not found. Please set it in your environment variables.")
    st.stop()

# Function to generate interview questions
def generate_questions(level, with_answers=False):
    prompt = f"""
    Generate 30 {level} interview questions for Data Science.
    {"Also provide answers." if with_answers else "Only provide questions."}
    """
    model = genai.GenerativeModel('gemini-1.5-flash')
    response = model.generate_content([prompt])
    return response.text if hasattr(response, 'text') else "No valid response."

def create_pdf(questions):
    """Generate a PDF with interview questions and return as BytesIO object."""
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    pdf.multi_cell(0, 10, questions)
    return io.BytesIO(pdf.output(dest="S").encode("latin1"))

def get_gemini_response(input_text, pdf_content, prompt):
    """Generate a response using Google Gemini API."""
    model = genai.GenerativeModel('gemini-1.5-flash')
    response = model.generate_content([input_text, pdf_content[0], prompt])
    return response.text if hasattr(response, 'text') else "No valid response from Gemini API"

def input_pdf_setup(uploaded_file):
    """Convert first page of uploaded PDF to an image and encode as base64."""
    if uploaded_file is not None:
        uploaded_file.seek(0)  # Reset file pointer
        images = pdf2image.convert_from_bytes(uploaded_file.read())  # Removed poppler_path
        first_page = images[0]
        img_byte_arr = io.BytesIO()
        first_page.save(img_byte_arr, format='JPEG')
        img_byte_arr = img_byte_arr.getvalue()
        pdf_parts = [{
            "mime_type": "image/jpeg",
            "data": base64.b64encode(img_byte_arr).decode()  # Encode to base64
        }]
        return pdf_parts
    else:
        raise FileNotFoundError("No File Uploaded")

def create_word_file(text):
    """Create a Word document with optimized resume content."""
    doc = Document()
    doc.add_heading("Updated Resume (Optimized for ATS)", level=1)
    doc.add_paragraph(text)
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def fetch_jobs(company):
    """Fetch jobs using JSearch API."""
    url = "https://jsearch.p.rapidapi.com/search"
    querystring = {"query": f"{company} Data Scientist", "num_pages": "1"}
    headers = {
        "X-RapidAPI-Key": JSEARCH_API_KEY,
        "X-RapidAPI-Host": "jsearch.p.rapidapi.com"
    }
    response = requests.get(url, headers=headers, params=querystring)
    if response.status_code == 200:
        jobs = response.json().get("data", [])
        return jobs
    else:
        return []

# Custom CSS for professional styling
st.markdown(
    """
    <style>
    .stButton button {
        background-color: #4CAF50;
        color: white;
        border-radius: 5px;
        padding: 10px 20px;
        font-size: 16px;
        border: none;
    }
    .stButton button:hover {
        background-color: #45a049;
    }
    .stTextArea textarea {
        font-size: 16px;
    }
    .stHeader {
        font-size: 36px;
        font-weight: bold;
        color: #4CAF50;
    }
    .stSubheader {
        font-size: 24px;
        font-weight: bold;
        color: #2E86C1;
    }
    .footer {
        position: fixed;
        left: 0;
        bottom: 0;
        width: 100%;
        background-color: #f1f1f1;
        text-align: center;
        padding: 10px;
    }
    </style>
    """,
    unsafe_allow_html=True,
)

# Streamlit App
st.markdown('<p class="stHeader">ATS Resume Expert & Interview Question Generator</p>', unsafe_allow_html=True)


# Prompts
input_prompt1 = """
You are an experienced HR with tech expertise in Data Science, Full Stack, Web Development, Big Data Engineering, DevOps, or Data Analysis.
Your task is to review the provided resume against the job description for these roles.
Please evaluate the candidate's profile, highlighting strengths and weaknesses in relation to the specified job role.
"""

input_prompt3 = """
You are a skilled ATS (Applicant Tracking System) scanner with expertise in Data Science, Full Stack, Web Development, Big Data Engineering, DevOps, and Data Analysis.
Your task is to evaluate the resume against the job description. Provide:
1. The percentage match.
2. Keywords missing.
3. Final evaluation.
"""

input_prompt4 = """
You are an experienced learning coach and technical expert. Create a 6-month personalized study plan for an individual aiming to excel in [Job Role], 
focusing on the skills, topics, and tools specified in the provided job description. Ensure the study plan includes:
- A list of topics and tools for each month.
- Suggested resources (books, online courses, documentation).
- Recommended practical exercises or projects.
- Periodic assessments or milestones.
- Tips for real-world applications.
"""

input_prompt5 = """
You are an AI-powered resume optimization expert. Your task is to enhance the provided resume for ATS (Applicant Tracking System) compatibility based on the job description.
Ensure the resume includes:
- Relevant keywords from the job description.
- Proper formatting for ATS readability.
- Optimized bullet points for skills and achievements.
- Improved professional summary.
- Updated job titles and descriptions as per industry standards.
Generate the optimized resume content in a structured format.
"""



# Sidebar for navigation
st.sidebar.title("Navigation")
app_mode = st.sidebar.radio("Choose the App Mode", ["Interview Question Generator", "ATS Resume Expert", "Job Search"])

if app_mode == "Interview Question Generator":
    st.markdown('<p class="stSubheader">📝 Generate Interview Questions</p>', unsafe_allow_html=True)
    topic = st.selectbox("Select Topic", ["Python", "Machine Learning", "Deep Learning", "Docker for Data Science"])
    level = st.radio("Select Level", ["Basic", "Intermediate", "Advanced"])
    if level:
        st.markdown(f'<p class="stSubheader">{level} Level Options</p>', unsafe_allow_html=True)
        option = st.radio("Choose Option", ["Only Questions", "With Question & Answer"])
        if st.button("Generate Questions"):
            with_answers = option == "With Question & Answer"
            questions = generate_questions(level, with_answers)
            st.text_area("Generated Questions", questions, height=300)
            pdf_file = create_pdf(questions)
            st.download_button(
                label="📥 Download Questions as PDF",
                data=pdf_file,
                file_name="interview_questions.pdf",
                mime="application/pdf"
            )

elif app_mode == "ATS Resume Expert":
    st.markdown('<p class="stSubheader">📄 ATS Resume Expert</p>', unsafe_allow_html=True)
    input_text = st.text_area("Job Description:", key="input", height=150)
    uploaded_file = st.file_uploader("Upload your resume (PDF)...", type=['pdf'])
    if uploaded_file:
        st.success("✅ PDF Uploaded Successfully.")
    col1, col2, col3, col4 = st.columns(4)
    with col1:
        submit1 = st.button("🔍 Tell Me About the Resume")
    with col2:
        submit3 = st.button("📊 Percentage Match")
    with col3:
        submit4 = st.button("📚 Personalized Learning Path")
    with col4:
        submit5 = st.button("✨ Optimize Resume & Download")
    if submit1:
        if uploaded_file:
            pdf_content = input_pdf_setup(uploaded_file)
            response = get_gemini_response(input_prompt1, pdf_content, input_text)
            st.subheader("The Response is:")
            st.write(response)
        else:
            st.warning("⚠️ Please upload a resume.")
    elif submit3:
        if uploaded_file:
            pdf_content = input_pdf_setup(uploaded_file)
            response = get_gemini_response(input_prompt3, pdf_content, input_text)
            st.subheader("The Response is:")
            st.write(response)
        else:
            st.warning("⚠️ Please upload a resume.")
    elif submit4:
        if uploaded_file:
            pdf_content = input_pdf_setup(uploaded_file)
            response = get_gemini_response(input_prompt4, pdf_content, input_text)
            st.subheader("The Response is:")
            st.write(response)
        else:
            st.warning("⚠️ Please upload a resume.")
    elif submit5:
        if uploaded_file:
            pdf_content = input_pdf_setup(uploaded_file)
            optimized_resume = get_gemini_response(input_prompt5, pdf_content, input_text)
            word_file = create_word_file(optimized_resume)
            st.subheader("Your optimized resume is ready to download:")
            st.download_button(
                label="📥 Download Optimized Resume (DOCX)",
                data=word_file,
                file_name="optimized_resume.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        else:
            st.warning("⚠️ Please upload a resume.")

elif app_mode == "Job Search":
    st.markdown('<p class="stSubheader">🔍 Explore Data Science & Analytics Jobs</p>', unsafe_allow_html=True)
    st.subheader("Click on a company to view job description:")
    companies = ["TCS", "Wipro", "Infosys", "Accenture", "Cognizant"]
    selected_company = None
    for company in companies:
        if st.button(company):
            selected_company = company
    if selected_company:
        st.subheader(f"Job Listings at {selected_company}")
        jobs = fetch_jobs(selected_company)
        if jobs:
            for job in jobs:
                st.markdown(f"### {job.get('job_title', 'Job Title Not Available')}")
                st.write(f"**Company:** {job.get('employer_name', 'N/A')}")
                st.write(f"**Location:** {job.get('job_city', 'Unknown')}, {job.get('job_country', 'Unknown')}")
                st.write(f"**Description:** {job.get('job_description', 'No description available.')}")
                st.markdown(f"[Apply Here]({job.get('job_apply_link', '#')})")
                st.write("---")
        else:
            st.write("No job listings found. Try again later!")

# Footer
st.markdown(
    '<div class="footer">Developed with ❤️ by Your Name</div>',
    unsafe_allow_html=True,
)