import streamlit as st
import os
import io
import base64
from PIL import Image
import pdf2image
import google.generativeai as genai
from docx import Document

# Load environment variables
from dotenv import load_dotenv
load_dotenv()

# Configure Google Gemini API
API_KEY = os.getenv("GOOGLE_API_KEY")
if not API_KEY:
    st.error("GOOGLE_API_KEY not found. Please set it in your environment variables.")
    st.stop()

genai.configure(api_key=API_KEY)

def get_gemini_response(input_text, pdf_content, prompt):
    """Generate a response using Google Gemini API."""
    model = genai.GenerativeModel('gemini-1.5-flash')
    response = model.generate_content([input_text, pdf_content[0], prompt])
    return response.text if hasattr(response, 'text') else "No valid response from Gemini API"

def input_pdf_setup(uploaded_file):
    """Convert first page of uploaded PDF to an image and encode as base64."""
    if uploaded_file is not None:
        uploaded_file.seek(0)  # Reset file pointer
        images = pdf2image.convert_from_bytes(uploaded_file.read())  # Removed poppler_path
        first_page = images[0]

        # Convert image to bytes
        img_byte_arr = io.BytesIO()
        first_page.save(img_byte_arr, format='JPEG')
        img_byte_arr = img_byte_arr.getvalue()

        pdf_parts = [{
            "mime_type": "image/jpeg",
            "data": base64.b64encode(img_byte_arr).decode()  # Encode to base64
        }]
        return pdf_parts
    else:
        raise FileNotFoundError("No File Uploaded")

def create_word_file(text):
    """Create a Word document with optimized resume content."""
    doc = Document()
    doc.add_heading("Updated Resume (Optimized for ATS)", level=1)
    doc.add_paragraph(text)
    
    # Save document in memory
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# Streamlit App
st.set_page_config(page_title="A5 ATS Resume Expert")
st.header("MY A5 PERSONAL ATS")

input_text = st.text_area("Job Description:", key="input")
uploaded_file = st.file_uploader("Upload your resume (PDF)...", type=['pdf'])

if uploaded_file:
    st.success("PDF Uploaded Successfully.")

# Buttons
submit1 = st.button("Tell Me About the Resume")
submit3 = st.button("Percentage Match")
submit4 = st.button("Personalized Learning Path")
submit5 = st.button("Optimize Resume & Download")

# Prompts
input_prompt1 = """
You are an experienced HR with tech expertise in Data Science, Full Stack, Web Development, Big Data Engineering, DevOps, or Data Analysis.
Your task is to review the provided resume against the job description for these roles.
Please evaluate the candidate's profile, highlighting strengths and weaknesses in relation to the specified job role.
"""

input_prompt3 = """
You are a skilled ATS (Applicant Tracking System) scanner with expertise in Data Science, Full Stack, Web Development, Big Data Engineering, DevOps, and Data Analysis.
Your task is to evaluate the resume against the job description. Provide:
1. The percentage match.
2. Keywords missing.
3. Final evaluation.
"""

input_prompt4 = """
You are an experienced learning coach and technical expert. Create a 6-month personalized study plan for an individual aiming to excel in [Job Role], 
focusing on the skills, topics, and tools specified in the provided job description. Ensure the study plan includes:
- A list of topics and tools for each month.
- Suggested resources (books, online courses, documentation).
- Recommended practical exercises or projects.
- Periodic assessments or milestones.
- Tips for real-world applications.
"""

input_prompt5 = """
You are an AI-powered resume optimization expert. Your task is to enhance the provided resume for ATS (Applicant Tracking System) compatibility based on the job description.
Ensure the resume includes:
- Relevant keywords from the job description.
- Proper formatting for ATS readability.
- Optimized bullet points for skills and achievements.
- Improved professional summary.
- Updated job titles and descriptions as per industry standards.
Generate the optimized resume content in a structured format.
"""

if submit1:
    if uploaded_file:
        pdf_content = input_pdf_setup(uploaded_file)
        response = get_gemini_response(input_prompt1, pdf_content, input_text)
        st.subheader("The Response is:")
        st.write(response)
    else:
        st.warning("Please upload a resume.")

elif submit3:
    if uploaded_file:
        pdf_content = input_pdf_setup(uploaded_file)
        response = get_gemini_response(input_prompt3, pdf_content, input_text)
        st.subheader("The Response is:")
        st.write(response)
    else:
        st.warning("Please upload a resume.")

elif submit4:
    if uploaded_file:
        pdf_content = input_pdf_setup(uploaded_file)
        response = get_gemini_response(input_prompt4, pdf_content, input_text)
        st.subheader("The Response is:")
        st.write(response)
    else:
        st.warning("Please upload a resume.")

elif submit5:
    if uploaded_file:
        pdf_content = input_pdf_setup(uploaded_file)
        optimized_resume = get_gemini_response(input_prompt5, pdf_content, input_text)
        
        # Create Word File
        word_file = create_word_file(optimized_resume)
        
        st.subheader("Your optimized resume is ready to download:")
        st.download_button(
            label="Download Optimized Resume (DOCX)",
            data=word_file,
            file_name="optimized_resume.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    else:
        st.warning("Please upload a resume.")
