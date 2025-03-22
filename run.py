import streamlit as st
import os
import io
import base64
from PIL import Image
import pdf2image
import google.generativeai as genai
from docx import Document
from fpdf import FPDF
from dotenv import load_dotenv
import requests
import speech_recognition as sr
from streamlit_mic_recorder import mic_recorder
from pydub import AudioSegment
import librosa
import numpy as np
import tempfile

# Set page config FIRST
st.set_page_config(page_title="ATS Resume & Interview Prep", layout="wide", initial_sidebar_state="expanded")

# Load environment variables
load_dotenv()

# Configure APIs
API_KEY = os.getenv("GOOGLE_API_KEY")
JSEARCH_API_KEY = os.getenv("JSEARCH_API_KEY")

if not API_KEY or not JSEARCH_API_KEY:
    st.error("API keys not found. Please set GOOGLE_API_KEY and JSEARCH_API_KEY in your environment variables.")
    st.stop()

genai.configure(api_key=API_KEY)

# Custom CSS for professional UI
st.markdown("""
    <style>
    .main { background-color: #f5f7fa; padding: 20px; }
    .stButton>button { background-color: #007bff; color: white; border-radius: 8px; padding: 10px 20px; font-size: 16px; border: none; }
    .stButton>button:hover { background-color: #0056b3; }
    .stTextArea>label, .stSelectbox>label, .stRadio>label { font-weight: bold; color: #333; }
    .stSubheader { color: #007bff; font-size: 24px; font-weight: bold; }
    .footer { position: fixed; left: 0; bottom: 0; width: 100%; background-color: #e9ecef; text-align: center; padding: 10px; font-size: 14px; color: #666; }
    .sidebar .sidebar-content { background-color: #ffffff; }
    </style>
""", unsafe_allow_html=True)

# Utility Functions
def generate_questions(level, with_answers=False):
    prompt = f"Generate 30 {level} interview questions for Data Science.\n{'Also provide answers.' if with_answers else 'Only provide questions.'}"
    model = genai.GenerativeModel('gemini-1.5-flash')
    response = model.generate_content([prompt])
    return response.text if hasattr(response, 'text') else "Error generating questions."

def create_pdf(text):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    pdf.multi_cell(0, 10, text)
    return io.BytesIO(pdf.output(dest="S").encode("latin1"))

def get_gemini_response(input_text, prompt="", pdf_content=None):
    model = genai.GenerativeModel('gemini-1.5-flash')
    content = [input_text, prompt] if not pdf_content else [input_text, pdf_content[0], prompt]
    response = model.generate_content(content)
    return response.text if hasattr(response, 'text') else "No valid response."

def input_pdf_setup(uploaded_file):
    uploaded_file.seek(0)
    images = pdf2image.convert_from_bytes(uploaded_file.read())
    first_page = images[0]
    img_byte_arr = io.BytesIO()
    first_page.save(img_byte_arr, format='JPEG')
    return [{"mime_type": "image/jpeg", "data": base64.b64encode(img_byte_arr.getvalue()).decode()}]

def create_word_file(text):
    doc = Document()
    doc.add_heading("Optimized Resume", level=1)
    doc.add_paragraph(text)
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def fetch_jobs(company):
    url = "https://jsearch.p.rapidapi.com/search"
    headers = {"X-RapidAPI-Key": JSEARCH_API_KEY, "X-RapidAPI-Host": "jsearch.p.rapidapi.com"}
    response = requests.get(url, headers=headers, params={"query": f"{company} Data Scientist", "num_pages": "1"})
    return response.json().get("data", []) if response.status_code == 200 else []

def evaluate_candidate(job_description, video_path):
    audio = AudioSegment.from_file(video_path, format="mp4")
    audio_path = video_path.replace(".mp4", ".wav")
    audio.export(audio_path, format="wav")
    y, sr = librosa.load(audio_path, sr=None)
    f0 = librosa.yin(y, fmin=50, fmax=400)
    valid_f0 = f0[f0 > 0]
    if len(valid_f0) == 0:
        return "No valid pitch detected."
    avg_pitch, pitch_variability = np.mean(valid_f0), np.std(valid_f0)
    confidence_score = max(0, min(100, 100 - (pitch_variability * 1.5)))
    prompt = f"""
    Evaluate candidate:
    Job Description: {job_description}
    Speech Analysis: Avg Pitch: {avg_pitch:.2f} Hz, Pitch Variability: {pitch_variability:.2f}, Confidence Score: {confidence_score:.1f}/100
    Assess projects from audio and alignment with job requirements.
    Provide: Qualification (Yes/No), Strengths/Weaknesses, Recommendations.
    """
    return get_gemini_response("", prompt)

def handle_voice_input():
    st.subheader("🎤 Voice Assistant")
    audio_dict = mic_recorder(start_prompt="🎤 Start Recording", stop_prompt="⏹ Stop", key="mic_recorder")
    if audio_dict and "bytes" in audio_dict:
        with st.spinner("Processing audio..."):
            audio_segment = AudioSegment.from_file(io.BytesIO(audio_dict["bytes"]), format="webm")
            wav_buffer = io.BytesIO()
            audio_segment.export(wav_buffer, format="wav")
            wav_buffer.seek(0)
            recognizer = sr.Recognizer()
            try:
                with sr.AudioFile(wav_buffer) as source:
                    audio = recognizer.record(source)
                    text = recognizer.recognize_google(audio)
                    st.text_area("Recognized Text", text, height=100)
                    response = get_gemini_response(text)
                    st.write("**Response:**", response)
            except sr.UnknownValueError:
                st.warning("Could not understand audio. Try again.")
            except Exception as e:
                st.error(f"Error: {e}")

# Mock Interview Functions
def generate_mock_question(level, topic):
    prompt = f"Generate a {level} difficulty interview question for {topic}."
    model = genai.GenerativeModel('gemini-1.5-flash')
    response = model.generate_content([prompt])
    return response.text if hasattr(response, 'text') else "Error generating question."

def evaluate_mock_answer(answer):
    prompt = f"Evaluate this interview answer: {answer}"
    model = genai.GenerativeModel('gemini-1.5-flash')
    response = model.generate_content([prompt])
    return response.text if hasattr(response, 'text') else "No valid response."

# Sidebar Navigation
st.sidebar.title("Navigation")
app_mode = st.sidebar.radio("Select Mode", [
    "Interview Questions", "ATS Resume Expert", "Job Search", "DSA Questions", "Candidate Evaluation", "Voice Assistant", "Mock Interview"
], index=0)

# Main App Logic
st.markdown('<h1 style="color: #007bff;">ATS Resume & Interview Prep</h1>', unsafe_allow_html=True)

if app_mode == "Interview Questions":
    st.subheader("Generate Interview Questions")
    level = st.selectbox("Level", ["Basic", "Intermediate", "Advanced"])
    with_answers = st.checkbox("Include Answers")
    if st.button("Generate"):
        with st.spinner("Generating..."):
            questions = generate_questions(level, with_answers)
            st.text_area("Questions", questions, height=300)
            pdf_file = create_pdf(questions)
            st.download_button("Download PDF", pdf_file, "questions.pdf", "application/pdf")

elif app_mode == "ATS Resume Expert":
    st.subheader("ATS Resume Expert")
    job_desc = st.text_area("Job Description", height=150)
    uploaded_file = st.file_uploader("Upload Resume (PDF)", type="pdf")
    if uploaded_file:
        col1, col2, col3, col4 = st.columns(4)
        with col1:
            if st.button("Analyze Resume"):
                pdf_content = input_pdf_setup(uploaded_file)
                response = get_gemini_response(job_desc, "Evaluate resume against job description.", pdf_content)
                st.write(response)
        with col2:
            if st.button("Match Percentage"):
                pdf_content = input_pdf_setup(uploaded_file)
                response = get_gemini_response(job_desc, "Calculate match percentage and missing keywords.", pdf_content)
                st.write(response)
        with col3:
            if st.button("Learning Path"):
                pdf_content = input_pdf_setup(uploaded_file)
                response = get_gemini_response(job_desc, "Create 6-month study plan.", pdf_content)
                st.write(response)
        with col4:
            if st.button("Optimize Resume"):
                pdf_content = input_pdf_setup(uploaded_file)
                optimized = get_gemini_response(job_desc, "Optimize resume for ATS.", pdf_content)
                word_file = create_word_file(optimized)
                st.download_button("Download Optimized Resume", word_file, "optimized_resume.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")

elif app_mode == "Job Search":
    st.subheader("Job Search")
    company = st.selectbox("Select Company", ["TCS", "Wipro", "Infosys", "Accenture", "Cognizant"])
    if st.button("Search Jobs"):
        jobs = fetch_jobs(company)
        for job in jobs:
            st.markdown(f"**{job.get('job_title')}** - {job.get('employer_name')}")
            st.write(f"Location: {job.get('job_city', 'N/A')}, {job.get('job_country', 'N/A')}")
            st.write(job.get('job_description', 'No description'))
            st.markdown(f"[Apply]({job.get('job_apply_link', '#')})")
            st.markdown("---")

elif app_mode == "DSA Questions":
    st.subheader("DSA Questions")
    level = st.selectbox("Difficulty", ["Easy", "Intermediate", "Advanced"])
    if st.button("Generate DSA Questions"):
        response = get_gemini_response(f"Generate {level} DSA questions for Data Science with answers.")
        st.write(response)

elif app_mode == "Candidate Evaluation":
    st.subheader("Candidate Evaluation")
    job_desc = st.text_area("Job Description", height=200)
    video = st.file_uploader("Upload Video (MP4)", type=["mp4"])
    if video and st.button("Evaluate"):
        with tempfile.NamedTemporaryFile(delete=False, suffix=".mp4") as tmp:
            tmp.write(video.read())
            evaluation = evaluate_candidate(job_desc, tmp.name)
            st.write(evaluation)
            os.remove(tmp.name)
            os.remove(tmp.name.replace(".mp4", ".wav"))

elif app_mode == "Voice Assistant":
    handle_voice_input()

elif app_mode == "Mock Interview":
    st.subheader("AI-Powered Mock Interview")
    topic = st.radio("Select Topic:", ("Python", "SQL"), index=0)
    level = st.radio("Select Difficulty:", ("Easy", "Intermediate", "Hard"), index=0)

    if 'mock_answers' not in st.session_state:
        st.session_state.mock_answers = []
    if 'mock_question' not in st.session_state:
        st.session_state.mock_question = None
    if 'mock_started' not in st.session_state:
        st.session_state.mock_started = False

    if st.button("Start Mock Interview"):
        st.session_state.mock_question = generate_mock_question(level, topic)
        st.session_state.mock_started = True
        st.rerun()

    if 'mock_started' in st.session_state and st.session_state.mock_started:
        st.write(f"**Question:** {st.session_state.mock_question}")
        
        if st.button("Click to Speak Your Answer"):
            recognizer = sr.Recognizer()
            with sr.Microphone() as source:
                st.info("Listening... Please speak your answer.")
                recognizer.adjust_for_ambient_noise(source)
                audio = recognizer.listen(source)
                try:
                    recognized_text = recognizer.recognize_google(audio)
                    st.text_area("Recognized Answer:", recognized_text, height=150)
                    
                    evaluation = evaluate_mock_answer(recognized_text)
                    st.subheader("Evaluation:")
                    st.write(evaluation)
                    
                    st.session_state.mock_answers.append(recognized_text)
                    st.session_state.mock_question = generate_mock_question(level, topic)
                    st.rerun()
                except sr.UnknownValueError:
                    st.warning("Could not understand the audio. Try again.")
                except sr.RequestError:
                    st.warning("Speech recognition service error.")
                except Exception as e:
                    st.warning(f"An error occurred: {e}")
        
        if len(st.session_state.mock_answers) >= 3:
            combined_answers = "\n".join(st.session_state.mock_answers)
            feedback = get_gemini_response(f"Overall feedback on these answers: {combined_answers}")
            st.subheader("Overall Feedback:")
            st.write(feedback)
            
            if st.button("Restart Mock Interview"):
                st.session_state.mock_answers = []
                st.session_state.mock_question = None
                st.session_state.mock_started = False
                st.rerun()

# Footer
st.markdown('<div class="footer">Developed with ❤️ by [Your Name]</div>', unsafe_allow_html=True)